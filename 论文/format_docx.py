# -*- coding: utf-8 -*-
"""
按苏州大学本科毕业论文格式要求生成Word文件
GB/T 7713.1-2006
"""
import re
import os
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# ========== 格式常量 ==========
FONT_SONG = '宋体'
FONT_HEI = '黑体'
FONT_TNR = 'Times New Roman'

# 字号映射（磅值）
SIZE_XIAO_ER = Pt(18)      # 小二号
SIZE_SAN_HAO = Pt(16)      # 三号
SIZE_XIAO_SAN = Pt(15)     # 小三号
SIZE_SI_HAO = Pt(14)       # 四号
SIZE_XIAO_SI = Pt(12)      # 小四号
SIZE_WU_HAO = Pt(10.5)     # 五号
SIZE_XIAO_WU = Pt(9)       # 小五号

# 页面设置（cm）
MARGIN_TOP = Cm(3.3)
MARGIN_BOTTOM = Cm(2.7)
MARGIN_LEFT = Cm(2.5)
MARGIN_RIGHT = Cm(2.5)
MARGIN_GUTTER = Cm(0.5)

HEADER_DIST = Cm(2.6)
FOOTER_DIST = Cm(2.0)


def set_cell_shading(cell, color):
    """设置单元格底色"""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def set_run_font(run, cn_font=FONT_SONG, en_font=FONT_TNR, size=SIZE_XIAO_SI, bold=False):
    """设置run的中英文字体"""
    run.font.size = size
    run.font.bold = bold
    run.font.name = en_font
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:eastAsia="{cn_font}"/>')
        rPr.insert(0, rFonts)
    else:
        rFonts.set(qn('w:eastAsia'), cn_font)


def set_paragraph_format(para, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                         line_spacing=1.5, space_before=Pt(0), space_after=Pt(0),
                         first_line_indent=None):
    """设置段落格式"""
    pf = para.paragraph_format
    pf.alignment = alignment
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = line_spacing
    pf.space_before = space_before
    pf.space_after = space_after
    if first_line_indent is not None:
        pf.first_line_indent = first_line_indent


def add_header(doc):
    """添加页眉：苏州大学本科生毕业设计（论文）"""
    section = doc.sections[0]
    header = section.header
    header.is_linked_to_previous = False
    para = header.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run('苏州大学本科生毕业设计（论文）')
    set_run_font(run, cn_font=FONT_SONG, size=SIZE_XIAO_WU, bold=False)
    # 添加页眉底部横线
    pPr = para._element.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:bottom w:val="single" w:sz="6" w:space="1" w:color="auto"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)


def setup_page(doc):
    """设置页面格式"""
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = MARGIN_TOP
    section.bottom_margin = MARGIN_BOTTOM
    section.left_margin = MARGIN_LEFT
    section.right_margin = MARGIN_RIGHT
    section.gutter = MARGIN_GUTTER
    section.header_distance = HEADER_DIST
    section.footer_distance = FOOTER_DIST


def add_title_page(doc):
    """添加标题页（封面简化版）"""
    # 空几行
    for _ in range(3):
        p = doc.add_paragraph()
        set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.CENTER)
        run = p.add_run('')
        set_run_font(run, size=SIZE_SAN_HAO)

    # 学校名称
    p = doc.add_paragraph()
    set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    run = p.add_run('苏州大学本科生毕业设计（论文）')
    set_run_font(run, cn_font=FONT_HEI, size=SIZE_XIAO_ER, bold=True)

    # 空行
    for _ in range(3):
        p = doc.add_paragraph()
        set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.CENTER)

    # 论文标题
    p = doc.add_paragraph()
    set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    run = p.add_run('基于手势识别的人车运动交互控制系统研究')
    set_run_font(run, cn_font=FONT_HEI, size=SIZE_XIAO_ER, bold=True)

    doc.add_page_break()


def add_section_title(doc, text, level=1):
    """添加章节标题"""
    p = doc.add_paragraph()
    if level == 1:
        # 章标题：小三号黑体加粗，居左
        set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.LEFT,
                             space_before=Pt(12), space_after=Pt(6))
        run = p.add_run(text)
        set_run_font(run, cn_font=FONT_HEI, en_font=FONT_TNR, size=SIZE_XIAO_SAN, bold=True)
    elif level == 2:
        # 节标题：小三号宋体加粗
        set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.LEFT,
                             space_before=Pt(6), space_after=Pt(3))
        run = p.add_run(text)
        set_run_font(run, cn_font=FONT_SONG, en_font=FONT_TNR, size=SIZE_XIAO_SAN, bold=True)
    elif level == 3:
        # 小节标题：四号宋体加粗
        set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.LEFT,
                             space_before=Pt(3), space_after=Pt(3))
        run = p.add_run(text)
        set_run_font(run, cn_font=FONT_SONG, en_font=FONT_TNR, size=SIZE_SI_HAO, bold=True)


def add_body_text(doc, text):
    """添加正文段落（小四号宋体，首行缩进2字符）"""
    p = doc.add_paragraph()
    set_paragraph_format(p, first_line_indent=Pt(24))  # 约2个字符
    # 分离中英文处理
    add_mixed_runs(p, text, size=SIZE_XIAO_SI, bold=False)
    return p


def add_mixed_runs(para, text, size=SIZE_XIAO_SI, bold=False):
    """处理中英文混排文本"""
    run = para.add_run(text)
    set_run_font(run, cn_font=FONT_SONG, en_font=FONT_TNR, size=size, bold=bold)


def add_table(doc, headers, rows, caption=None):
    """添加表格"""
    if caption:
        p = doc.add_paragraph()
        set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                             space_before=Pt(6), space_after=Pt(3))
        run = p.add_run(caption)
        set_run_font(run, cn_font=FONT_SONG, size=SIZE_WU_HAO, bold=True)

    num_cols = len(headers)
    num_rows = len(rows) + 1  # +1 for header
    table = doc.add_table(rows=num_rows, cols=num_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    # 表头
    for j, header in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(header)
        set_run_font(run, cn_font=FONT_SONG, size=SIZE_WU_HAO, bold=True)
        set_cell_shading(cell, "D9E2F3")

    # 数据行
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = table.rows[i + 1].cells[j]
            cell.text = ''
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(str(val))
            set_run_font(run, cn_font=FONT_SONG, size=SIZE_WU_HAO, bold=False)

    # 表后空一行
    doc.add_paragraph()
    return table


def parse_md_and_build_doc(md_path, output_path):
    """解析Markdown并生成格式化Word"""
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    doc = Document()
    setup_page(doc)
    add_header(doc)

    # ===== 封面页 =====
    add_title_page(doc)

    # ===== 中文摘要 =====
    p = doc.add_paragraph()
    set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=Pt(12), space_after=Pt(6))
    run = p.add_run('摘要')
    set_run_font(run, cn_font=FONT_HEI, size=SIZE_XIAO_ER, bold=True)

    # ===== 解析正文 =====
    i = 0
    skip_to_chapter1 = True  # 跳过md开头的标题和目录部分

    while i < len(lines):
        line = lines[i].rstrip('\n')
        stripped = line.strip()

        # 跳过空行
        if not stripped:
            i += 1
            continue

        # 跳过封面信息（md文件开头部分）
        if skip_to_chapter1:
            if stripped.startswith('第1章'):
                skip_to_chapter1 = False
                # 不 i+=1, 继续处理这一行
            else:
                # 处理摘要部分
                if stripped == '摘要':
                    i += 1
                    continue
                elif stripped.startswith('关键词'):
                    p = doc.add_paragraph()
                    set_paragraph_format(p, first_line_indent=Pt(24))
                    run = p.add_run('关键词：')
                    set_run_font(run, cn_font=FONT_HEI, size=SIZE_XIAO_SI, bold=True)
                    kw_text = stripped.replace('关键词：', '').replace('关键词:', '')
                    run2 = p.add_run(kw_text)
                    set_run_font(run2, cn_font=FONT_SONG, size=SIZE_XIAO_SI)
                    i += 1
                    continue
                elif stripped == 'ABSTRACT' or stripped.startswith('ABSTRACT'):
                    # 英文摘要页
                    doc.add_page_break()
                    p = doc.add_paragraph()
                    set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=Pt(12), space_after=Pt(6))
                    run = p.add_run('ABSTRACT')
                    set_run_font(run, en_font=FONT_TNR, size=SIZE_XIAO_ER, bold=True)
                    i += 1
                    continue
                elif stripped.startswith('Keywords'):
                    p = doc.add_paragraph()
                    set_paragraph_format(p, first_line_indent=Pt(24))
                    run = p.add_run('Keywords: ')
                    set_run_font(run, en_font=FONT_TNR, size=SIZE_XIAO_SI, bold=True)
                    kw_text = stripped.replace('Keywords:', '').replace('Keywords：', '').strip()
                    run2 = p.add_run(kw_text)
                    set_run_font(run2, en_font=FONT_TNR, size=SIZE_XIAO_SI)
                    i += 1
                    continue
                elif stripped.startswith('Research on'):
                    # 英文标题
                    doc.add_page_break()
                    p = doc.add_paragraph()
                    set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=Pt(24))
                    run = p.add_run(stripped)
                    set_run_font(run, en_font=FONT_TNR, size=SIZE_XIAO_ER, bold=True)
                    i += 1
                    continue
                elif stripped == '目录':
                    # 目录页（占位，Word中可自动生成）
                    doc.add_page_break()
                    p = doc.add_paragraph()
                    set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=Pt(12))
                    run = p.add_run('目录')
                    set_run_font(run, cn_font=FONT_HEI, size=SIZE_XIAO_ER, bold=True)
                    p2 = doc.add_paragraph()
                    set_paragraph_format(p2)
                    run2 = p2.add_run('（请在Word中插入自动目录：引用 → 目录 → 自动目录）')
                    set_run_font(run2, cn_font=FONT_SONG, size=SIZE_XIAO_SI)
                    doc.add_page_break()
                    i += 1
                    continue
                elif stripped.startswith('苏州大学') or stripped.startswith('基于手势识别'):
                    i += 1
                    continue
                elif stripped.startswith('第') and '章' not in stripped:
                    # 可能是目录条目，跳过
                    i += 1
                    continue
                elif re.match(r'^\d+\.\d+', stripped) or re.match(r'^\d+\.', stripped):
                    # 目录子条目，跳过
                    i += 1
                    continue
                elif stripped.startswith('总结与展望') or stripped.startswith('参考文献') or stripped.startswith('致谢'):
                    i += 1
                    continue
                elif stripped.startswith('  '):
                    # 缩进的目录条目，跳过
                    i += 1
                    continue
                else:
                    # 摘要正文
                    if stripped and not stripped.startswith('本文的主要'):
                        add_body_text(doc, stripped)
                    elif stripped.startswith('本文的主要'):
                        add_body_text(doc, stripped)
                    i += 1
                    continue

        # ===== 章标题 =====
        if re.match(r'^第\d+章\s', stripped):
            # 新章节需要分页（第2章开始）
            if not stripped.startswith('第1章'):
                doc.add_page_break()
            add_section_title(doc, stripped, level=1)
            i += 1
            continue

        # ===== 总结与展望 =====
        if stripped == '总结与展望':
            doc.add_page_break()
            add_section_title(doc, stripped, level=1)
            i += 1
            continue

        # ===== 参考文献 =====
        if stripped == '参考文献':
            doc.add_page_break()
            add_section_title(doc, stripped, level=1)
            i += 1
            continue

        # ===== 致谢 =====
        if stripped == '致谢':
            doc.add_page_break()
            add_section_title(doc, stripped, level=1)
            i += 1
            continue

        # ===== 参考文献条目 =====
        if stripped.startswith('[') and ']' in stripped[:5]:
            p = doc.add_paragraph()
            set_paragraph_format(p)
            run = p.add_run(stripped)
            set_run_font(run, cn_font=FONT_SONG, en_font=FONT_TNR, size=SIZE_WU_HAO)
            i += 1
            continue

        # ===== 节标题 (X.X) =====
        if re.match(r'^\d+\.\d+\s', stripped):
            add_section_title(doc, stripped, level=2)
            i += 1
            continue

        # ===== 小节标题 (X.X.X) =====
        if re.match(r'^\d+\.\d+\.\d+\s', stripped):
            add_section_title(doc, stripped, level=3)
            i += 1
            continue

        # ===== 表格处理 =====
        if stripped.startswith('表') and re.match(r'^表\d+-?\d*', stripped):
            caption = stripped
            i += 1
            # 跳过表格标题和表格内容之间的空行
            while i < len(lines) and not lines[i].strip().startswith('|'):
                i += 1
            # 收集表格内容
            table_lines = []
            while i < len(lines):
                tl = lines[i].strip()
                if tl.startswith('|') and '|' in tl[1:]:
                    table_lines.append(tl)
                    i += 1
                else:
                    break
            # 解析表格
            if len(table_lines) >= 2:
                headers = [c.strip() for c in table_lines[0].split('|')[1:-1]]
                # 跳过分隔行
                data_lines = [tl for tl in table_lines[1:] if not re.match(r'^[\|:\-\s]+$', tl)]
                rows = []
                for dl in data_lines:
                    cells = [c.strip().replace('**', '') for c in dl.split('|')[1:-1]]
                    rows.append(cells)
                add_table(doc, headers, rows, caption=caption)
            continue

        # ===== 普通段落 =====
        if stripped.startswith('$$') or stripped.startswith('$'):
            # 公式（简单处理）
            p = doc.add_paragraph()
            set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.CENTER)
            formula = stripped.replace('$$', '').replace('$', '')
            run = p.add_run(formula)
            set_run_font(run, en_font=FONT_TNR, size=SIZE_XIAO_SI)
            i += 1
            continue

        if stripped.startswith('（此处'):
            # 图片占位符
            p = doc.add_paragraph()
            set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.CENTER)
            run = p.add_run('【' + stripped + '】')
            set_run_font(run, cn_font=FONT_SONG, size=SIZE_WU_HAO)
            i += 1
            continue

        if stripped.startswith('图') and (stripped[1:2].isdigit() or stripped.startswith('图1-') or stripped.startswith('图2-') or stripped.startswith('图3-') or stripped.startswith('图4-')):
            # 图标题
            p = doc.add_paragraph()
            set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.CENTER)
            run = p.add_run(stripped)
            set_run_font(run, cn_font=FONT_SONG, size=SIZE_WU_HAO)
            i += 1
            continue

        # 普通正文
        add_body_text(doc, stripped)
        i += 1

    doc.save(output_path)
    print(f'论文已保存至: {output_path}')


if __name__ == '__main__':
    md_path = r'D:\workspace\GestureBot\论文\论文_完整版.md'
    output_path = r'D:\workspace\GestureBot\论文\论文_格式化版.docx'
    parse_md_and_build_doc(md_path, output_path)
